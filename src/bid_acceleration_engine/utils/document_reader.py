"""Universal document reader supporting .txt, .docx, and .pdf formats."""

from pathlib import Path

import pdfplumber
from docx import Document


def read_document(file_path: Path) -> str:
    """Read and extract text from .txt, .docx, or .pdf files.

    Args:
        file_path: Path to document file (.txt, .docx, or .pdf)

    Returns:
        Extracted text content as a single string

    Raises:
        ValueError: If file format is not supported
        FileNotFoundError: If file does not exist
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"Document not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return _read_text_file(file_path)
    elif suffix == ".docx":
        return _read_docx_file(file_path)
    elif suffix == ".pdf":
        return _read_pdf_file(file_path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}. Supported: .txt, .docx, .pdf")


def _read_text_file(file_path: Path) -> str:
    """Read plain text file."""
    with open(file_path, encoding="utf-8") as f:
        return f.read()


def _read_docx_file(file_path: Path) -> str:
    """Extract text from Word document.

    Preserves structure by:
    - Extracting paragraph text in order
    - Including table content
    - Maintaining section breaks as newlines
    """
    doc = Document(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables (preserve structure)
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                text_parts.append("\t".join(row_cells))

    return "\n".join(text_parts)


def _read_pdf_file(file_path: Path) -> str:
    """Extract text from PDF file.

    Handles:
    - Multi-page documents
    - Text extraction from pages
    - Table content (best-effort)
    - Preserves page structure via newlines
    """
    text_parts = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # Extract text from page
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(page_text)

            # Try to extract table data if present
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        row_text = [str(cell).strip() if cell else "" for cell in row]
                        row_text = [cell for cell in row_text if cell]  # Filter empty cells
                        if row_text:
                            text_parts.append("\t".join(row_text))

            # Add page break marker for multi-page docs
            if page_num < len(pdf.pages):
                text_parts.append("\n" + "=" * 80 + "\n")

    return "\n".join(text_parts)
